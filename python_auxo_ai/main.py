from flask import *
import PyPDF2
import docx
import os
import zipfile
from html_write import *

import pandas as pd

write_to_html()

app = Flask(__name__,template_folder='D:\\python_auxo_ai\\template\\')
  
  
@app.route('/')
def main():
    return render_template("index.html")
  
@app.route('/upload', methods=['POST'])
def upload():

    if request.method == 'POST':
  
        # Get the list of files from webpage
        files = request.files.getlist("file")
        file_n=[]
        
        output_dir = "D:\python_auxo_ai\output_folder"  
        output_docx="D:\python_auxo_ai\output_docx"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        def bookmark_dict(bookmark_list):
            result = {}
            for item in bookmark_list:
                
                if isinstance(item, list):
                    # recursive call
                    result.update(bookmark_dict(item))
                else:
                    result[reader.get_destination_page_number(item)] = item.title
            return result


        def get_content_dict(res,filename):
            pdfReader = PyPDF2.PdfReader(filename)
            content={}
            for keys in res:
                
                
            # Creating a page object
                pageObj = pdfReader.pages[int(keys)]
                
                # Extracting text from page
                data = pageObj.extract_text()
                content[keys]=data
            return  content
                # Closing the pdf file object

        # converting to docx
        def convert_to_docx(res,content,filename):
            data=[]
            doc = docx.Document()
            for keys in res:
                doc.add_heading("Main Title :   " + res[keys])
                doc_para = doc.add_paragraph('Content:')
                
                doc_para.add_run(content[keys].strip())
            
            # add a page break to start a new page
                doc.add_page_break()
                mydict = { "Section":res[keys], "data" : content[keys]}
                data.append(mydict)
            df = pd.DataFrame.from_dict(data)
            print(df)
            doc.save(os.path.join(output_docx, filename))
            csv_file=os.path.splitext(filename)[0]+".csv"
            df.to_csv(os.path.join(output_docx, csv_file))

        for file in files:
                file_path = os.path.join(output_dir, file.filename)
                file.save(file_path)
                print(file.filename)
                reader = PyPDF2.PdfReader(file.filename)
                res=bookmark_dict(reader.outline) # pg no and heading
                
                content=get_content_dict(res,file.filename) 
                 # pg no and content

                file_name=file.filename
                docx_file = os.path.splitext(file_name)[0] + ".docx"
                convert_to_docx(res,content,docx_file)  
                file_n.append(file.filename)

                # convert_to_docx(res,content,file_name)  
                print(file_n)
                
        zip_path = "D:\python_auxo_ai\zipfile.zip"  # Replace with the path where you want to save the zip file
        with zipfile.ZipFile(zip_path, 'w') as zip_file:
            for filename in os.listdir(output_docx):
                print(filename)
                file_path = os.path.join(output_docx, filename)
                zip_file.write(file_path, arcname=filename)
                os.unlink(output_docx+"\\" + filename)




        #download_link = f'<a href="/download"><button>Download</button></a>'
        return send_file(zip_path, as_attachment=True,attachment_filename='converted_files.zip')  
    
    # Send the zip file as an attachment to the client
    #return send_file(zip_path, as_attachment=True, attachment_filename='converted_files.zip')
        #return render_template("output.html", names = file_n) 

    
  
if __name__ == '__main__':
    app.run(debug=True)
