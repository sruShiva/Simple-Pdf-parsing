import PyPDF2
import docx

def bookmark_dict(bookmark_list):
    result = {}
    for item in bookmark_list:
        
        if isinstance(item, list):
            # recursive call
            result.update(bookmark_dict(item))
        else:
            result[reader.get_destination_page_number(item)] = item.title
    return result


def get_content_dict(res):
    pdfReader = PyPDF2.PdfReader("D:\python_auxo_ai\BCBS_Plan1.pdf")
    content={}
    for keys in res:
        
       
    # Creating a page object
        pageObj = pdfReader.pages[int(keys)]
        
        # Extracting text from page
        data = pageObj.extract_text()
        content[keys]=data
    return  content
        # Closing the pdf file object

reader = PyPDF2.PdfReader("D:\python_auxo_ai\BCBS_Plan1.pdf")
res=bookmark_dict(reader.outline) # pg no and heading
content=get_content_dict(res)  # pg no and content

file_name="D:\python_auxo_ai\BCBS_Plan1.pdf"
# converting to docx
def convert_to_docx(res,content,filename):
    doc = docx.Document()
    for keys in res:
        doc.add_heading("Main Title :   " + res[keys])
        doc_para = doc.add_paragraph('Content:')
        
        doc_para.add_run(content[keys].strip())
    
    # add a page break to start a new page
        doc.add_page_break()
        doc.save(filename+".docx")

convert_to_docx(res,content,file_name)

# zip_path = "/path/to/zip/archive.zip"  # Replace with the path where you want to save the zip file
# with zipfile.ZipFile(zip_path, 'w') as zip_file:
#     for file_name in file_list:
#             file_path = os.path.join(files_dir, file_name)
#             zip_file.write(file_path, arcname=file_name)

#     # Send the zip file as an attachment
# return send_file(zip_path, as_attachment=True, attachment_filename='files.zip')
